# -*- coding: utf-8 -*-
"""Generate Aircraft Training Content Player user manual (.docx)"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

DOCS_DIR = os.path.dirname(__file__)
IMAGES_DIR = os.path.join(DOCS_DIR, "images")
OUTPUT = os.path.join(DOCS_DIR, "Aircraft_Training_Content_Player_UserManual.docx")

FIGURE_IMAGES = {
    1: "fig01_main_screen.png",
    2: "fig02_bin_folder.png",
    3: "fig03_startup.png",
    4: "fig04_layout.png",
    5: "fig05_course_tree.png",
    6: "fig06_lesson_selected.png",
    7: "fig07_video_playback.png",
    8: "fig08_pdf_view.png",
    9: "fig09_image_view.png",
    10: "fig10_complete.png",
    11: "fig11_workflow.png",
    12: "fig12_refresh.png",
    13: "fig13_json_edit.png",
    14: "fig14_add_lesson.png",
    15: "fig15_bin_folder.png",
}


def set_doc_defaults(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "맑은 고딕"
    font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    for i in range(1, 4):
        hs = doc.styles[f"Heading {i}"]
        hs.font.name = "맑은 고딕"
        hs._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        hs.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\nAircraft Training Content Player\n")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("사용자 매뉴얼\n")
    r2.bold = True
    r2.font.size = Pt(22)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run("항공정비 교육용 콘텐츠 통합 재생 프로그램").font.size = Pt(14)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("문서 버전: 1.0\n")
    info.add_run("작성일: 2026년 6월\n")
    info.add_run("대상 독자: 교수 · 학생 · 교육기관 관리자\n")

    doc.add_page_break()


def add_toc(doc):
    doc.add_heading("목  차", level=1)
    toc_items = [
        "1. 개요",
        "2. 시스템 요구사항",
        "3. 설치 방법",
        "4. 프로그램 실행",
        "5. 메인 화면 설명",
        "6. 과정(Course) 선택 방법",
        "7. 강의(Lesson) 선택 방법",
        "8. 동영상 보기 기능",
        "9. PDF 보기 기능",
        "10. 이미지 보기 기능",
        "11. 학습 완료 및 진행 상태 관리",
        "12. Course JSON 파일 구조 설명",
        "13. 새로운 강의 추가 방법",
        "14. 자주 발생하는 문제 해결",
        "15. FAQ",
        "부록 A. 폴더 구조 참고",
        "부록 B. 용어 설명",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Number")
    doc.add_page_break()


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_figure(doc, num, caption):
    image_name = FIGURE_IMAGES.get(num)
    image_path = os.path.join(IMAGES_DIR, image_name) if image_name else None

    if image_path and os.path.isfile(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run().add_picture(image_path, width=Inches(6.0))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[그림 {num}]")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(f"그림 {num}. {caption}")
    cap_run.italic = True
    cap_run.font.size = Pt(10)
    doc.add_paragraph()


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def build_manual():
    doc = Document()
    set_doc_defaults(doc)

    # margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_title_page(doc)
    add_toc(doc)

    # 1. 개요
    add_heading(doc, "1. 개요")
    add_para(doc,
        "Aircraft Training Content Player(이하 '본 프로그램')는 항공정비 교육을 위한 "
        "통합 학습 콘텐츠 재생 프로그램입니다. YouTube 교육 영상, PDF 교재, 이미지 자료를 "
        "하나의 화면에서 선택하여 학습할 수 있도록 설계되었습니다.")
    add_para(doc,
        "본 프로그램은 외부 서버에 의존하지 않고, 실행 파일과 함께 배포되는 Data 폴더의 "
        "JSON 설정 파일을 읽어 강의 목록을 구성합니다. 따라서 교육기관에서는 JSON 파일과 "
        "학습 자료 파일만 교체·추가하면 새로운 과정을 손쉽게 반영할 수 있습니다.")
    add_heading(doc, "1.1 목적", level=2)
    add_bullets(doc, [
        "항공정비 관련 이론·실습 교육 콘텐츠의 통합 제공",
        "YouTube 영상, PDF 문서, 이미지 자료의 일원화된 접근",
        "강의별 학습 완료 상태 기록 및 관리",
        "교육기관별 맞춤 과정 구성(코스·레슨 확장)",
    ])
    add_heading(doc, "1.2 대상 독자", level=2)
    add_table(doc, ["역할", "주요 사용 목적"], [
        ["교수", "강의 진행, 동영상·PDF·이미지 자료 시연, 학습 완료 확인"],
        ["학생", "선택한 강의의 영상 시청, PDF 열람, 이미지 확인, 학습 완료 표시"],
        ["교육기관 관리자", "JSON·PDF·이미지 자료 배포, 신규 과정 추가, 프로그램 설치"],
    ])
    add_heading(doc, "1.3 주요 기능 요약", level=2)
    add_table(doc, ["기능", "설명"], [
        ["과정 트리 탐색", "좌측 트리에서 Course(과정)와 Lesson(강의) 선택"],
        ["영상보기", "YouTube 교육 영상을 프로그램 내 콘텐츠 영역에서 재생"],
        ["PDF보기", "연결된 PDF 파일을 기본 PDF 뷰어로 열기"],
        ["이미지보기", "이미지 확대 보기 창에서 학습 이미지 확인"],
        ["학습완료", "선택한 강의의 완료 상태를 Progress.json에 저장"],
        ["새로고침", "Data 폴더의 JSON 변경 사항을 다시 불러오기"],
        ["이전 / 다음", "같은 과정 내 이전·다음 강의로 이동"],
    ])
    add_figure(doc, 1, "Aircraft Training Content Player 전체 화면 구성")

    # 2. 시스템 요구사항
    add_heading(doc, "2. 시스템 요구사항")
    add_heading(doc, "2.1 하드웨어 요구사항", level=2)
    add_table(doc, ["항목", "최소 사양", "권장 사양"], [
        ["운영체제", "Windows 10 (64비트)", "Windows 10/11 (64비트)"],
        ["프로세서", "Intel Core i3 이상", "Intel Core i5 이상"],
        ["메모리", "4 GB RAM", "8 GB RAM 이상"],
        ["디스플레이", "1280 × 720", "1920 × 1080 이상"],
        ["저장 공간", "500 MB 여유 공간", "2 GB 이상(자료 포함)"],
        ["네트워크", "YouTube 영상 재생 시 인터넷 필요", "안정적인 브로드밴드 연결"],
    ])
    add_heading(doc, "2.2 소프트웨어 요구사항", level=2)
    add_bullets(doc, [
        "Microsoft Edge WebView2 Runtime (영상보기 기능에 필수)",
        "PDF 열람용 프로그램 (Microsoft Edge, Adobe Acrobat Reader 등)",
        "인터넷 연결 (YouTube 영상 스트리밍 시)",
    ])
    add_para(doc,
        "WebView2 Runtime은 Microsoft에서 무료로 제공합니다. 설치되어 있지 않은 경우 "
        "영상보기 실행 시 오류 메시지가 표시되며, 아래 주소에서 Evergreen Bootstrapper를 "
        "다운로드하여 설치할 수 있습니다.")
    add_para(doc,
        "▶ https://developer.microsoft.com/microsoft-edge/webview2/")
    add_heading(doc, "2.3 지원 파일 형식", level=2)
    add_table(doc, ["자료 유형", "지원 형식", "비고"], [
        ["과정 설정", "JSON (UTF-8)", "Data 폴더에 저장"],
        ["PDF", ".pdf", "Pdf 폴더에 저장"],
        ["이미지", ".png, .jpg, .bmp 등", "Images 폴더에 저장"],
        ["학습 진행", "JSON (UTF-8)", "Progress\\Progress.json"],
    ])

    # 3. 설치 방법
    add_heading(doc, "3. 설치 방법")
    add_para(doc,
        "본 프로그램은 별도의 설치 마법사 없이 폴더 복사 방식으로 배포·설치합니다. "
        "교육기관 PC 또는 학습용 PC에 아래 폴더 구조 전체를 복사하면 설치가 완료됩니다.")
    add_heading(doc, "3.1 배포 폴더 구조", level=2)
    add_para(doc, "기본 설치 경로 예시:")
    add_para(doc, "C:\\TrainingCenter\\Bin\\", bold=True)
    add_para(doc, "폴더 구성:")
    code_lines = [
        "Bin\\",
        "  ├── TrainingContentPlayer.exe    ← 실행 파일",
        "  ├── Data\\                        ← 과정 JSON 파일",
        "  │     ├── AircraftHydraulic.json",
        "  │     ├── AircraftElectric.json",
        "  │     ├── AircraftAerodynamics.json",
        "  │     └── (기타 과정 JSON)",
        "  ├── Pdf\\                         ← PDF 학습 자료",
        "  ├── Images\\                      ← 이미지 학습 자료",
        "  └── Progress\\                    ← 학습 진행 기록",
        "        └── Progress.json",
    ]
    for line in code_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Cm(1)
    add_figure(doc, 2, "Bin 폴더 구조 (탐색기 화면)")
    add_heading(doc, "3.2 설치 절차", level=2)
    steps = [
        "관리자가 제공한 Bin 폴더 전체를 대상 PC의 원하는 위치에 복사합니다.",
        "TrainingContentPlayer.exe 파일이 Data, Pdf, Images, Progress 하위 폴더와 "
        "같은 Bin 폴더 안에 있는지 확인합니다.",
        "WebView2 Runtime이 설치되어 있는지 확인합니다. (미설치 시 2.2절 참고)",
        "바탕화면에 TrainingContentPlayer.exe 바로 가기를 만들면 편리합니다.",
        "최초 실행 전 백신·보안 프로그램에서 실행 파일 차단 여부를 확인합니다.",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")
    add_heading(doc, "3.3 업데이트 방법", level=2)
    add_bullets(doc, [
        "실행 파일만 교체: TrainingContentPlayer.exe를 새 버전으로 덮어씁니다.",
        "교육 자료만 교체: Data, Pdf, Images 폴더의 파일을 추가·수정합니다.",
        "자료 변경 후 프로그램에서 「새로고침」 버튼을 클릭하여 반영합니다.",
        "Progress.json은 학습 기록이므로 업데이트 시 삭제하지 않는 것을 권장합니다.",
    ])

    # 4. 프로그램 실행
    add_heading(doc, "4. 프로그램 실행")
    add_heading(doc, "4.1 실행 방법", level=2)
    steps = [
        "Bin 폴더로 이동합니다.",
        "TrainingContentPlayer.exe 파일을 더블 클릭합니다.",
        "잠시 후 메인 화면(Training Content Player)이 표시됩니다.",
        "좌측 트리에 등록된 과정 목록이 자동으로 로드됩니다.",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")
    add_figure(doc, 3, "프로그램 실행 직후 메인 화면")
    add_heading(doc, "4.2 프로그램 종료", level=2)
    add_bullets(doc, [
        "메인 창 오른쪽 상단의 「X」(닫기) 버튼을 클릭합니다.",
        "종료 시 학습 진행 기록(Progress.json)은 자동으로 유지됩니다.",
        "영상 재생 중 종료하면 재생이 자동으로 중지됩니다.",
    ])
    add_heading(doc, "4.3 창 크기 조절", level=2)
    add_para(doc,
        "메인 창은 최소화, 최대화, 크기 조절이 가능합니다. 창 크기를 변경하면 "
        "좌측 트리, 설명 영역, 콘텐츠 영역이 자동으로 재배치됩니다. "
        "동영상 재생 중에도 콘텐츠 영역 크기에 맞게 영상 화면이 조정됩니다.")

    doc.add_page_break()

    # 5. 메인 화면 설명
    add_heading(doc, "5. 메인 화면 설명")
    add_para(doc,
        "메인 화면은 크게 네 영역으로 구성됩니다. 각 영역의 이름과 역할을 이해하면 "
        "프로그램을 보다 효율적으로 사용할 수 있습니다.")
    add_figure(doc, 4, "메인 화면 영역 구분 (①~⑥)")
    add_table(doc, ["번호", "영역", "설명"], [
        ["①", "제목 표시줄", "프로그램 이름(Training Content Player), 최소화·최대화·닫기 버튼"],
        ["②", "과정 트리 (좌측)", "Course(과정)와 Lesson(강의) 목록. 클릭하여 강의 선택"],
        ["③", "강의 제목", "현재 선택된 Lesson의 제목 표시"],
        ["④", "강의 설명", "선택된 Lesson의 상세 설명 (최소 2줄 표시)"],
        ["⑤", "기능 버튼", "영상보기, PDF보기, 이미지보기, 새로고침, 다음, 이전"],
        ["⑥", "콘텐츠 영역", "썸네일 이미지 또는 동영상 재생 화면"],
        ["⑦", "학습완료 버튼", "화면 하단 오른쪽. 현재 강의 완료 처리"],
    ])
    add_heading(doc, "5.1 기능 버튼 설명", level=2)
    add_table(doc, ["버튼 이름", "기능", "활성 조건"], [
        ["영상보기", "YouTube 교육 영상을 콘텐츠 영역에서 재생", "해당 강의에 YoutubeUrl이 등록된 경우"],
        ["PDF보기", "연결된 PDF 파일을 외부 뷰어로 열기", "해당 강의에 PdfFile이 등록된 경우"],
        ["이미지보기", "이미지 확대 보기 창에서 학습 이미지 표시", "해당 강의에 ImageFile이 등록된 경우"],
        ["새로고침", "Data 폴더의 JSON 파일을 다시 읽어 목록 갱신", "항상 사용 가능"],
        ["다음", "같은 과정 또는 다음 과정의 다음 강의로 이동", "다음 강의가 존재할 때"],
        ["이전", "이전 강의로 이동", "이전 강의가 존재할 때"],
        ["학습완료", "현재 강의를 완료 처리하고 Progress.json에 저장", "강의가 선택된 경우"],
    ])
    add_para(doc,
        "※ 버튼이 비활성(회색)으로 표시되면 해당 강의에 연결된 자료가 없거나, "
        "이동할 이전·다음 강의가 없는 상태입니다.")

    # 6. Course 선택
    add_heading(doc, "6. 과정(Course) 선택 방법")
    add_para(doc,
        "Course(과정)는 항공 유압, 항공 전기, 항공역학 등 교육 주제 단위의 상위 그룹입니다. "
        "Data 폴더의 JSON 파일 하나가 하나의 Course에 해당합니다.")
    add_heading(doc, "6.1 과정 목록 확인", level=2)
    steps = [
        "프로그램을 실행하면 좌측 트리에 Course 이름이 최상위 항목으로 표시됩니다.",
        "Course 이름 왼쪽의 ▶ 또는 ▼ 아이콘으로 하위 Lesson 목록을 펼치거나 접을 수 있습니다.",
        "기본 제공 과정 예: Aircraft Hydraulic System, Aircraft Electrical System, "
        "Aircraft Aerodynamics, Drone Operation, Japanese Language",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")
    add_figure(doc, 5, "좌측 트리에서 Course 목록 확인")
    add_heading(doc, "6.2 과정 펼치기/접기", level=2)
    add_bullets(doc, [
        "Course 이름 왼쪽 ▶ 아이콘 클릭 → 하위 Lesson 목록이 펼쳐집니다.",
        "▼ 아이콘 클릭 → Lesson 목록이 접힙니다.",
        "프로그램 최초 실행 시 과정은 기본적으로 펼쳐진 상태로 표시됩니다.",
    ])

    # 7. Lesson 선택
    add_heading(doc, "7. 강의(Lesson) 선택 방법")
    add_para(doc,
        "Lesson(강의)은 Course 안의 개별 학습 단원입니다. Lesson을 선택하면 "
        "제목, 설명, 썸네일 이미지가 오른쪽 화면에 표시됩니다.")
    add_heading(doc, "7.1 강의 선택 절차", level=2)
    steps = [
        "좌측 트리에서 원하는 Course를 펼칩니다.",
        "학습할 Lesson 이름을 마우스로 클릭합니다.",
        "오른쪽 상단에 Lesson 제목이 표시됩니다.",
        "제목 아래 설명 영역에 강의 개요가 표시됩니다.",
        "콘텐츠 영역에 해당 Lesson의 썸네일 이미지가 표시됩니다.",
        "영상보기, PDF보기, 이미지보기 버튼이 해당 자료 유무에 따라 활성화됩니다.",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")
    add_figure(doc, 6, "Lesson 선택 후 화면 변화")
    add_heading(doc, "7.2 학습 완료 표시 확인", level=2)
    add_para(doc,
        "이미 「학습완료」 처리한 Lesson은 트리 목록에서 [완료] 접두어가 붙어 표시됩니다. "
        "예: [완료] Introduction")
    add_heading(doc, "7.3 이전·다음 강의 이동", level=2)
    add_bullets(doc, [
        "「다음」 버튼: 현재 Course의 다음 Lesson으로 이동. 마지막 Lesson이면 다음 Course의 첫 Lesson으로 이동",
        "「이전」 버튼: 현재 Course의 이전 Lesson으로 이동. 첫 Lesson이면 이전 Course의 마지막 Lesson으로 이동",
        "첫 번째·마지막 강의에서는 해당 방향 버튼이 비활성화됩니다",
    ])

    doc.add_page_break()

    # 8. 동영상 보기
    add_heading(doc, "8. 동영상 보기 기능")
    add_para(doc,
        "「영상보기」 버튼을 클릭하면 선택한 Lesson에 연결된 YouTube 교육 영상이 "
        "프로그램 내 콘텐츠 영역에서 재생됩니다. 별도의 웹 브라우저 창이 열리지 않습니다.")
    add_heading(doc, "8.1 영상 재생 방법", level=2)
    steps = [
        "학습할 Lesson을 선택합니다.",
        "「영상보기」 버튼이 활성화되어 있는지 확인합니다.",
        "「영상보기」 버튼을 클릭합니다.",
        "콘텐츠 영역에 YouTube 영상이 재생됩니다.",
        "영상 아래 또는 위에 YouTube 기본 재생 컨트롤(일시정지, 음량 등)을 사용할 수 있습니다.",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")
    add_figure(doc, 7, "콘텐츠 영역에서 YouTube 영상 재생 화면")
    add_heading(doc, "8.2 영상 전환 및 중지", level=2)
    add_bullets(doc, [
        "다른 Lesson을 선택하면 이전 영상은 자동으로 중지되고 썸네일이 표시됩니다.",
        "다른 Lesson에서 다시 「영상보기」를 클릭하면 새 영상이 같은 영역에서 재생됩니다.",
        "「PDF보기」 또는 「이미지보기」 클릭 시에도 재생 중인 영상이 중지됩니다.",
    ])
    add_heading(doc, "8.3 영상 재생 관련 주의사항", level=2)
    add_bullets(doc, [
        "YouTube 영상 재생에는 인터넷 연결이 필요합니다.",
        "일부 영상은 업로더의 설정에 따라 임베드 재생이 제한될 수 있습니다.",
        "영상이 재생되지 않으면 JSON의 YoutubeUrl이 올바른지 확인하세요. (12·13·14절 참고)",
    ])

    # 9. PDF 보기
    add_heading(doc, "9. PDF 보기 기능")
    add_para(doc,
        "「PDF보기」 버튼을 클릭하면 해당 Lesson에 연결된 PDF 파일이 "
        "Windows 기본 PDF 연결 프로그램(Edge, Adobe Reader 등)으로 열립니다.")
    add_heading(doc, "9.1 PDF 열람 방법", level=2)
    steps = [
        "PDF 자료가 있는 Lesson을 선택합니다.",
        "「PDF보기」 버튼을 클릭합니다.",
        "Pdf 폴더에 저장된 PDF 파일이 외부 뷰어에서 열립니다.",
        "PDF 열람 후 학습 프로그램으로 돌아와 계속 학습할 수 있습니다.",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")
    add_figure(doc, 8, "PDF보기 실행 — 외부 PDF 뷰어 화면")
    add_heading(doc, "9.2 PDF 파일을 찾을 수 없는 경우", level=2)
    add_para(doc,
        "「파일을 찾을 수 없습니다」 메시지가 표시되면 JSON의 PdfFile 경로와 "
        "실제 Pdf 폴더 내 파일 존재 여부를 확인하세요.")

    # 10. 이미지 보기
    add_heading(doc, "10. 이미지 보기 기능")
    add_para(doc,
        "「이미지보기」 버튼을 클릭하면 별도의 이미지 보기 창이 열리고, "
        "해당 Lesson의 학습 이미지를 확대하여 볼 수 있습니다.")
    add_heading(doc, "10.1 이미지 보기 방법", level=2)
    steps = [
        "이미지 자료가 있는 Lesson을 선택합니다.",
        "「이미지보기」 버튼을 클릭합니다.",
        "이미지 보기 창에 학습 이미지가 표시됩니다.",
        "확인 후 이미지 보기 창을 닫고 메인 화면으로 돌아옵니다.",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")
    add_figure(doc, 9, "이미지 보기 창")
    add_para(doc,
        "※ 콘텐츠 영역에 표시되는 썸네일은 Lesson 선택 시 자동으로 표시되는 미리보기 이미지이며, "
        "「이미지보기」는 별도의 확대 보기 창을 제공합니다.")

    # 11. 학습 완료
    add_heading(doc, "11. 학습 완료 및 진행 상태 관리")
    add_heading(doc, "11.1 학습완료 처리", level=2)
    steps = [
        "학습을 마친 Lesson을 선택합니다.",
        "화면 하단 오른쪽의 「학습완료」 버튼을 클릭합니다.",
        "「학습 완료 상태가 저장되었습니다」 메시지가 표시됩니다.",
        "좌측 트리에서 해당 Lesson 앞에 [완료] 표시가 추가됩니다.",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")
    add_figure(doc, 10, "학습완료 처리 후 트리의 [완료] 표시")
    add_heading(doc, "11.2 진행 기록 파일", level=2)
    add_para(doc,
        "학습 완료 정보는 Bin\\Progress\\Progress.json 파일에 저장됩니다. "
        "프로그램을 다시 실행해도 완료 상태가 유지됩니다.")
    add_para(doc, "Progress.json 구조 예시:")
    example = (
        '{\n'
        '  "CompletedLessons": [\n'
        '    { "CourseName": "Aircraft Hydraulic System", "LessonTitle": "Introduction" },\n'
        '    { "CourseName": "Aircraft Hydraulic System", "LessonTitle": "Reservoir" }\n'
        '  ]\n'
        '}'
    )
    p = doc.add_paragraph(example)
    p.paragraph_format.left_indent = Cm(1)

    doc.add_page_break()

    # 12. JSON 구조
    add_heading(doc, "12. Course JSON 파일 구조 설명")
    add_para(doc,
        "각 Course는 Data 폴더에 있는 하나의 JSON 파일로 정의됩니다. "
        "파일 인코딩은 UTF-8이어야 하며, 한글·영문 모두 정상 표시됩니다.")
    add_heading(doc, "12.1 JSON 전체 구조", level=2)
    json_example = (
        '{\n'
        '  "CourseName": "Aircraft Hydraulic System",\n'
        '  "Lessons": [\n'
        '    {\n'
        '      "Title": "Introduction",\n'
        '      "Description": "강의 설명 (여러 줄 가능)",\n'
        '      "YoutubeUrl": "https://www.youtube.com/watch?v=VIDEO_ID",\n'
        '      "PdfFile": "Pdf\\\\Intro.pdf",\n'
        '      "ImageFile": "Images\\\\Intro.png"\n'
        '    }\n'
        '  ]\n'
        '}'
    )
    p = doc.add_paragraph(json_example)
    p.paragraph_format.left_indent = Cm(1)
    add_figure(doc, 13, "JSON 파일 편집 예시 (메모장 또는 Visual Studio Code)")
    add_heading(doc, "12.2 필드 설명", level=2)
    add_table(doc, ["필드명", "위치", "필수", "설명"], [
        ["CourseName", "최상위", "예", "과정 이름. 트리 최상위에 표시"],
        ["Lessons", "최상위", "예", "강의 배열"],
        ["Title", "Lesson", "예", "강의 제목. 트리 하위 항목에 표시"],
        ["Description", "Lesson", "예", "강의 설명. \\n으로 줄바꿈 가능"],
        ["YoutubeUrl", "Lesson", "아니오", "YouTube watch URL. 없으면 영상보기 비활성"],
        ["PdfFile", "Lesson", "아니오", "PDF 상대 경로. 없으면 PDF보기 비활성"],
        ["ImageFile", "Lesson", "아니오", "이미지 상대 경로. 없으면 이미지보기 비활성"],
    ])
    add_heading(doc, "12.3 경로 작성 규칙", level=2)
    add_bullets(doc, [
        "PdfFile, ImageFile은 Bin 폴더 기준 상대 경로입니다.",
        "경로 구분자는 백슬래시(\\)를 사용합니다. JSON에서는 \\\\로 기록합니다.",
        "예: \"PdfFile\": \"Pdf\\\\Intro.pdf\"",
        "예: \"ImageFile\": \"Images\\\\Reservoir.png\"",
        "YoutubeUrl은 https://www.youtube.com/watch?v=영상ID 형식을 권장합니다.",
        "YoutubeUrl, PdfFile, ImageFile 중 사용하지 않는 항목은 빈 문자열(\"\")로 둡니다.",
    ])

    # 13. 새 강의 추가
    add_heading(doc, "13. 새로운 강의 추가 방법")
    add_heading(doc, "13.1 기존 과정에 Lesson 추가", level=2)
    steps = [
        "Data 폴더에서 해당 Course의 JSON 파일을 텍스트 편집기로 엽니다.",
        "Lessons 배열에 새 Lesson 객체를 추가합니다.",
        "Title, Description, YoutubeUrl, PdfFile, ImageFile을 작성합니다.",
        "해당 PDF·이미지 파일을 Pdf, Images 폴더에 복사합니다.",
        "JSON 파일을 UTF-8로 저장합니다.",
        "프로그램에서 「새로고침」 버튼을 클릭하여 변경 사항을 반영합니다.",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")
    add_figure(doc, 14, "JSON 파일에 Lesson 추가하는 화면")
    add_heading(doc, "13.2 새로운 Course 추가", level=2)
    steps = [
        "Data 폴더에 새 JSON 파일을 만듭니다. (예: AircraftLandingGear.json)",
        "CourseName과 Lessons 배열을 작성합니다.",
        "관련 Pdf, Images 파일을 각 폴더에 배치합니다.",
        "프로그램에서 「새로고침」을 클릭하면 좌측 트리에 새 Course가 표시됩니다.",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")
    add_heading(doc, "13.3 YouTube URL 확인 방법 (관리자용)", level=2)
    add_para(doc,
        "잘못된 YouTube URL을 등록하면 「동영상을 재생할 수 없음」 오류가 발생합니다. "
        "새 URL 등록 전 아래를 확인하세요.")
    add_bullets(doc, [
        "브라우저에서 해당 URL이 정상 재생되는지 확인",
        "삭제되었거나 비공개인 영상은 사용할 수 없음",
        "watch?v=영상ID 형식의 URL 사용 권장",
        "교육용 공개 영상(TED-Ed, MIT OCW, FAA 등) 사용 권장",
    ])

    # 학습 시나리오
    add_heading(doc, "11.3 교수·학생 학습 시나리오", level=2)
    add_heading(doc, "교수용 시나리오", level=3)
    prof_steps = [
        "수업 시작 전 Bin 폴더의 Data·Pdf·Images 자료가 최신인지 확인합니다.",
        "프로그램을 실행하고 좌측 트리에서 당일 수업 Lesson을 선택합니다.",
        "콘텐츠 영역의 썸네일으로 수업 주제를 간략히 소개합니다.",
        "「영상보기」로 YouTube 교육 영상을 콘텐츠 영역에서 재생합니다.",
        "영상 시청 후 「PDF보기」로 교재를 열어 이론을 보충 설명합니다.",
        "필요 시 「이미지보기」로 구조·회로·부품 이미지를 확대하여 설명합니다.",
        "수업 마무리 시 학생들이 「학습완료」를 눌러 진행 상태를 기록하도록 안내합니다.",
    ]
    for i, step in enumerate(prof_steps, 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")
    add_figure(doc, 11, "교수 수업 진행 흐름 (Lesson 선택 → 영상 → PDF → 학습완료)")
    add_heading(doc, "학생용 시나리오", level=3)
    student_steps = [
        "프로그램을 실행하고 배정된 Course를 트리에서 찾습니다.",
        "학습할 Lesson을 클릭하여 제목·설명을 읽습니다.",
        "「영상보기」로 강의 영상을 시청합니다.",
        "「PDF보기」로 교재를 열어 필기·복습합니다.",
        "「이미지보기」로 그림 자료를 확대 확인합니다.",
        "「다음」 버튼으로 다음 Lesson으로 이동하며 순서대로 학습합니다.",
        "학습을 마치면 「학습완료」를 클릭하여 완료 상태를 저장합니다.",
    ]
    for i, step in enumerate(student_steps, 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")

    add_heading(doc, "11.4 새로고침 기능", level=2)
    add_para(doc,
        "관리자가 Data 폴더의 JSON 파일을 수정·추가한 경우, 실행 중인 프로그램에서는 "
        "「새로고침」 버튼을 클릭해야 변경 사항이 반영됩니다. 새로고침 시 현재 선택 중이던 "
        "강의 제목을 기준으로 동일 Lesson을 다시 찾아 선택합니다.")
    add_figure(doc, 12, "「새로고침」 버튼 위치")

    doc.add_page_break()

    # 14. 문제 해결
    add_heading(doc, "14. 자주 발생하는 문제 해결")
    problems = [
        ("프로그램이 실행되지 않음",
         ["Windows 64비트 환경인지 확인",
          "백신 프로그램의 차단 목록에서 실행 파일 허용",
          "관리자 권한으로 실행 시도"]),
        ("좌측 트리에 과정이 표시되지 않음",
         ["Bin\\Data 폴더에 JSON 파일이 있는지 확인",
          "JSON 파일이 UTF-8 인코딩인지 확인",
          "JSON 문법 오류(쉼표, 괄호) 여부 확인",
          "「새로고침」 버튼 클릭"]),
        ("WebView2 초기화 오류 (영상보기 불가)",
         ["Microsoft Edge WebView2 Runtime 설치",
          "인터넷 연결 확인",
          "프로그램 재시작"]),
        ("동영상을 재생할 수 없음",
         ["YoutubeUrl이 올바른 공개 영상인지 확인",
          "인터넷 연결 상태 확인",
          "JSON 수정 후 「새로고침」 실행"]),
        ("PDF/이미지 파일을 찾을 수 없음",
         ["PdfFile, ImageFile 경로가 Bin 기준 상대 경로인지 확인",
          "실제 파일이 Pdf, Images 폴더에 존재하는지 확인",
          "JSON에서 백슬래시가 \\\\로 이스케이프되었는지 확인"]),
        ("한글이 깨져서 표시됨",
         ["JSON 파일을 UTF-8로 저장",
          "메모장 저장 시 인코딩을 UTF-8로 선택",
          "Visual Studio Code 사용 시 우측 하단 UTF-8 확인"]),
        ("학습완료가 저장되지 않음",
         ["Bin\\Progress 폴더 쓰기 권한 확인",
          "Progress.json 파일이 읽기 전용이 아닌지 확인"]),
        ("프로그램 종료 시 오류 발생",
         ["프로그램을 최신 버전으로 업데이트",
          "영상 재생 중이면 중지 후 종료",
          "WebView2 Runtime 재설치"]),
    ]
    for title, solutions in problems:
        add_heading(doc, f"14.{problems.index((title, solutions)) + 1} {title}", level=2)
        add_bullets(doc, solutions)

    # 15. FAQ
    add_heading(doc, "15. FAQ (자주 묻는 질문)")
    faqs = [
        ("Q1. 인터넷 없이 사용할 수 있나요?",
         "PDF와 이미지는 오프라인에서 사용 가능합니다. YouTube 영상은 인터넷 연결이 필요합니다."),
        ("Q2. 여러 JSON 파일을 동시에 사용할 수 있나요?",
         "예. Data 폴더의 모든 JSON 파일이 자동으로 로드되어 트리에 표시됩니다."),
        ("Q3. JSON을 수정했는데 반영이 안 됩니다.",
         "프로그램에서 「새로고침」 버튼을 클릭하세요. 실행 중인 프로그램은 자동으로 파일 변경을 감지하지 않습니다."),
        ("Q4. 영상이 브라우저가 아닌 프로그램 안에서 재생되나요?",
         "예. 「영상보기」는 프로그램 내 콘텐츠 영역에서 재생됩니다."),
        ("Q5. 학습 완료 기록을 초기화하려면?",
         "Bin\\Progress\\Progress.json 파일에서 CompletedLessons 배열을 비우거나 파일을 삭제 후 재실행하세요."),
        ("Q6. 영상·PDF·이미지 중 일부만 있는 강의도 만들 수 있나요?",
         "예. 없는 자료의 URL/경로를 빈 문자열(\"\")로 두면 해당 버튼이 비활성화됩니다."),
        ("Q7. Course 표시 순서는 어떻게 결정되나요?",
         "Data 폴더의 JSON 파일 로드 순서에 따릅니다. 특정 순서가 필요하면 파일 이름을 알파벳 순으로 조정하세요."),
        ("Q8. Mac/Linux에서 사용할 수 있나요?",
         "아니요. 본 프로그램은 Windows 64비트 전용입니다."),
        ("Q9. YouTube 외의 영상 URL을 사용할 수 있나요?",
         "현재 버전은 YouTube URL만 지원합니다."),
        ("Q10. 교육기관 로고나 과정명을 변경할 수 있나요?",
         "과정명은 JSON의 CourseName으로 변경합니다. 프로그램 제목 창 이름은 Training Content Player로 고정되어 있습니다."),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        rq = p.add_run(q + "\n")
        rq.bold = True
        doc.add_paragraph(a)

    # Appendix A
    doc.add_page_break()
    add_heading(doc, "부록 A. 폴더 구조 참고")
    add_table(doc, ["폴더/파일", "역할", "편집 주체"], [
        ["TrainingContentPlayer.exe", "프로그램 실행 파일", "개발자"],
        ["Data\\*.json", "과정·강의 정의", "관리자"],
        ["Pdf\\*.pdf", "PDF 교재", "관리자"],
        ["Images\\*.png", "이미지·썸네일", "관리자"],
        ["Progress\\Progress.json", "학습 완료 기록", "프로그램 자동 생성"],
    ])
    add_figure(doc, 15, "교육기관 PC에 배포된 Bin 폴더 전체 구조")

    add_heading(doc, "13.4 JSON 작성 실전 예제", level=2)
    add_para(doc, "아래는 Aircraft Aerodynamics 과정에 Lesson 하나를 추가하는 예제입니다.")
    add_example = (
        '    {\n'
        '      "Title": "Wing Structure",\n'
        '      "Description": "Wing Structure Overview\\n\\n'
        '날개 구조의 기본 구성 요소와 역할을 학습합니다.",\n'
        '      "YoutubeUrl": "https://www.youtube.com/watch?v=p4VHMsIuPmk",\n'
        '      "PdfFile": "Pdf\\\\AeroWingStructure.pdf",\n'
        '      "ImageFile": "Images\\\\AeroWingStructure.png"\n'
        '    }'
    )
    p = doc.add_paragraph(add_example)
    p.paragraph_format.left_indent = Cm(1)
    add_para(doc,
        "위 항목을 Lessons 배열의 마지막 항목 뒤에 쉼표(,)로 구분하여 추가합니다. "
        "Pdf, Images 폴더에 실제 파일을 배치한 후 「새로고침」을 실행하세요.")

    add_heading(doc, "13.5 관리자 점검 체크리스트", level=2)
    checklist = [
        "JSON 파일이 UTF-8 인코딩으로 저장되었는가?",
        "CourseName이 다른 Course와 중복되지 않는가?",
        "각 Lesson의 Title이 같은 Course 내에서 고유한가?",
        "YoutubeUrl이 브라우저에서 정상 재생되는가?",
        "PdfFile, ImageFile 경로에 실제 파일이 존재하는가?",
        "JSON 문법 검사(괄호, 쉼표, 따옴표)를 완료했는가?",
        "프로그램 「새로고침」 후 트리에 정상 표시되는가?",
    ]
    for item in checklist:
        doc.add_paragraph(f"☐ {item}")

    add_heading(doc, "5.2 화면 크기별 사용 팁", level=2)
    add_bullets(doc, [
        "노트북(1366×768): 프로그램을 최대화하여 콘텐츠 영역을 넓게 사용하세요.",
        "대형 모니터(1920×1080 이상): 창을 최대화하면 트리·설명·영상이 여유 있게 표시됩니다.",
        "프로젝터 연결 시: 영상 재생 전 「영상보기」를 눌러 콘텐츠 영역에 영상을 띄운 후 확대하세요.",
    ])

    # Appendix B
    add_heading(doc, "부록 B. 용어 설명")
    add_table(doc, ["용어", "설명"], [
        ["Course (과정)", "교육 주제 단위. JSON 파일 하나에 해당"],
        ["Lesson (강의)", "과정 내 개별 학습 단원"],
        ["Data 폴더", "Course JSON 파일이 저장되는 폴더"],
        ["WebView2", "Microsoft Edge 기반 내장 웹 엔진. 영상 재생에 사용"],
        ["Progress.json", "학습 완료 상태를 저장하는 파일"],
        ["새로고침", "JSON 데이터를 다시 읽어 화면을 갱신하는 기능"],
        ["콘텐츠 영역", "화면 오른쪽의 썸네일·동영상 표시 영역"],
    ])

    # footer note
    doc.add_paragraph()
    add_para(doc,
        "— 문서 끝 —\n"
        "Aircraft Training Content Player 사용자 매뉴얼 v1.0\n"
        "문의: 교육기관 시스템 관리자")

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    build_manual()
